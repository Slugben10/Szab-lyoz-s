import os
import json
import requests
from flask import Flask, request, jsonify, render_template
from bs4 import BeautifulSoup
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import CharacterTextSplitter
import openai
from docx import Document as DocxDocument
from langchain.schema import Document
from langdetect import detect  # Import the langdetect library

# OpenAI API key setup
os.environ['OPENAI_API_KEY'] = ''

# Load JSON rules
with open('chatbot_rules.json', 'r') as file:
    rules = json.load(file)['rules']  # Directly access the list of rules

# Setup Flask app
app = Flask(__name__)

# OpenAI client configuration
openai.api_key = os.getenv('OPENAI_API_KEY')

# Initialize vector database
persist_directory = '/Users/binobenjamin/Documents/Szabályozás/chroma_data'

# Function to load DOCX documents
def load_docx(file_path):
    doc = DocxDocument(file_path)
    return [Document(page_content=para.text, metadata={'source': file_path}) for para in doc.paragraphs if para.text]

# Load and split documents from the specified directory
documents_folder = '/Users/binobenjamin/Documents/Szabályozás/Dokumentumok'
all_documents = []

# Load documents from folder
for filename in os.listdir(documents_folder):
    if filename.endswith(".txt"):
        loader = TextLoader(os.path.join(documents_folder, filename))
        all_documents.extend(loader.load())
    elif filename.endswith(".docx"):
        docx_documents = load_docx(os.path.join(documents_folder, filename))
        all_documents.extend(docx_documents)
    elif filename.endswith(".pdf"):
        pdf_loader = PyPDFLoader(os.path.join(documents_folder, filename))
        all_documents.extend(pdf_loader.load())

# Check if documents are loaded correctly
print("Loaded documents:", all_documents)

# Flatten the list to get only 'page_content'
documents_to_split = [doc for doc in all_documents if isinstance(doc, Document)]

# Split documents into chunks
text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
documents = text_splitter.split_documents(documents_to_split)

# Store documents in vector database
db = Chroma.from_documents(
    documents,
    OpenAIEmbeddings(),
    persist_directory=persist_directory
)

# Web scraping function
def scrape_and_store(url):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            paragraphs = [p.get_text() for p in soup.find_all('p')]
            scraped_text = ' '.join(paragraphs)[:1000]

            # Split the scraped text into chunks
            scraped_documents = text_splitter.split_documents([Document(page_content=scraped_text)])
            
            # Store scraped documents in vector database
            db.add_documents(scraped_documents)
            return scraped_text
        else:
            return "Error: Unable to access the website."
    except Exception as e:
        return f"Error scraping the website: {str(e)}"

# Helper function to format rules into numbered points
def format_rules():
    formatted_rules = []
    for rule in rules:
        if rule['active']:
            # Check if the rule has a description and convert to numbered points if it's long
            if rule['description']:
                formatted_description = rule['description'].replace("\n", " ").strip()
                points = formatted_description.split('. ')
                formatted_rules.append(f"{rule['name']}:\n")
                for i, point in enumerate(points, start=1):
                    formatted_rules.append(f"{i}. {point.strip()}")
                formatted_rules.append("\n")
    return "\n".join(formatted_rules)

# Function to translate response based on language
def translate_response(response, language):
    if language != 'en':
        try:
            # Using a translation API or a translation library like googletrans
            from googletrans import Translator
            translator = Translator()
            translated = translator.translate(response, src='en', dest=language)
            return translated.text
        except Exception as e:
            print(f"Error in translation: {e}")
            return response
    return response

# Serve HTML frontend
@app.route('/')
def home():
    return render_template('index.html')

# Chat route to process user messages and call the GPT-3.5-turbo API
@app.route('/chat', methods=['POST'])
def chat():
    user_input = request.json.get('message')
    url_to_scrape = request.json.get('url')

    if not user_input:
        return jsonify({'response': "Error: No input provided"}), 400

    scraped_data = ''
    if url_to_scrape:
        scraped_data = scrape_and_store(url_to_scrape)

    try:
        # Detect user input language
        language = detect(user_input)

        # Search for matching document in vector database
        docs = db.similarity_search(user_input)
        if docs:
            document_match = docs[0].page_content
        else:
            document_match = "No matching document found."

        # System message for API call
        initial_message = "Hello! How can I assist you today?"
        max_response_length = 300
        specific_phrases = ""

        system_message = (
            f"{initial_message}\n\n"
            f"You are a helpful assistant. Below is information scraped from a website and a related document that might be helpful:\n\n"
            f"Web scraped data: {scraped_data}\n\n"
            f"Relevant document: {document_match}\n\n"
            f"Specific phrases to consider: {specific_phrases}\n\n"
            f"Here are some guidelines you should follow:\n\n"
            f"{format_rules()}\n\n"
            f"Now, respond to the user's question in detail, using no more than {max_response_length} characters."
        )

        # Make an API call with the system message
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_input},
            ]
        )
        bot_response = response['choices'][0]['message']['content']

        # Translate the response if necessary
        translated_response = translate_response(bot_response, language)

        return jsonify({'response': translated_response})

    except Exception as e:
        return jsonify({'response': f"Error: {str(e)}"})

if __name__ == "__main__":
    app.run(debug=True)
